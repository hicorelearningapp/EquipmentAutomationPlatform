import sys
import subprocess

def install_and_import():
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    
    import docx
    return docx

docx = install_and_import()
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_compliance_sheet():
    doc_path = r'e:\Github\EquipmentAutomationPlatforms\GEM_Compliance_Replacement.docx'
    print(f"Creating compliance sheet at: {doc_path}")
    
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Title / Heading
    h = doc.add_heading("8. GEM Compliance Verification", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p_intro = doc.add_paragraph(
        "This section defines the GEM compliance level for the equipment based on SEMI E30 standards. "
        "The compliance sheet tables below outline both the fundamental requirements and additional capabilities "
        "supported by the equipment interface."
    )
    p_intro.paragraph_format.space_after = Pt(12)
    
    # Data definition
    fundamental_reqs = [
        ("State Models", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Equipment Processing States", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Host-Initiated S1F13/14 Scenario", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Event Notification", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Online Identification", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Error Messages", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Documentation", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Control (Operator Initiated)", "☑ Yes ☐ No", "☑ Yes ☐ No"),
    ]
    
    additional_caps = [
        ("Establish Communications", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Dynamic Event Report Configuration", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Data Variable and Collection Event Namelist", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Variable Data Collection", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Trace Data Collection", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Status Data Collection", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Alarm Management", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Remote Control", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Equipment Constants", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Process Program Management", "☑ Yes ☐ No", "☑ Yes ☐ No (See note E30-04)"),
        ("Material Movement", "☐ Yes ☑ No", "☐ Yes ☑ No (See note E30-05)"),
        ("Equipment Terminal Services", "☑ Yes ☐ No", "☑ Yes ☐ No"),
        ("Clock", "☑ Yes ☐ No", "☑ Yes ☐ No (See note E30-06)"),
        ("Limits Monitoring", "☐ Yes ☑ No", "☐ Yes ☑ No (See note E30-07)"),
        ("Spooling", "☑ Yes ☐ No", "☑ Yes ☐ No (See note E30-08)"),
        ("Control (Host-Initiated)", "☑ Yes ☐ No", "☑ Yes ☐ No"),
    ]
    
    def add_compliance_table(title, data):
        # Create Table: Rows = len(data) + 1 (header), Cols = 3
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Format Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = title
        hdr_cells[1].text = "Implemented"
        hdr_cells[2].text = "GEM Compliant"
        
        # Apply bold to header row
        for cell in hdr_cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        # Add Data Rows
        for req, impl, comp in data:
            row_cells = table.add_row().cells
            row_cells[0].text = req
            row_cells[1].text = impl
            row_cells[2].text = comp
            
            # Formatting formatting for columns
            for i, cell in enumerate(row_cells):
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(10)
                        
        # Add space after table
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(12)

    # Add both tables
    add_compliance_table("Fundamental GEM Requirements", fundamental_reqs)
    add_compliance_table("Additional Capabilities", additional_caps)
    
    # Add Notes section
    h_notes = doc.add_heading("Notes", level=2)
    h_notes.paragraph_format.space_before = Pt(12)
    h_notes.paragraph_format.space_after = Pt(6)
    
    notes = [
        "E30-01 – Equipment accepts incoming multi-block messages but does not require Inquiry/Grant.",
        "E30-02 – Equipment never sends S9F13 (Conversation Timeout); effectively an infinite timeout.",
        "E30-03 – Host-initiated S1F13/S1F14 communication supported.",
        "E30-04 – Process Program Management supports unformatted recipes.",
        "E30-05 – Material Movement support is limited; full E87 carrier management is not claimed.",
        "E30-06 – Clock synchronization via S2F17/S2F18.",
        "E30-07 – Limits monitoring capability not implemented.",
        "E30-08 – Spooling supported using S2F43/S2F44 and S6F23/S6F24."
    ]
    
    for note in notes:
        p_note = doc.add_paragraph(note)
        p_note.paragraph_format.left_indent = Inches(0.25)
        p_note.paragraph_format.space_after = Pt(3)
        for run in p_note.runs:
            run.font.size = Pt(9.5)
            
    doc.save(doc_path)
    print("Document generated successfully!")

if __name__ == "__main__":
    create_compliance_sheet()
