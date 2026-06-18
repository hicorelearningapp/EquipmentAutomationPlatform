import sys
import subprocess

def install_and_import():
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    
    import docx
    return docx

docx = install_and_import()
from docx import Document

# The list of items to append
items_to_append = [
    "SSEC SECS GEM – TOOL NAME – SEC GEM MANUAL",
    "(search for smls - knowldedge)",
    "(add format column in svid, dvid)",
    "Add equipment constants table",
    "M,",
    "EQ ID",
    "STANDARDS SUPPORTED",
    "Gem COMPLIANCE SHEET",
    "HSMS CONFIGURATION – IMP: DEVIDE ID, ip address, timeouts, baud rate",
    "Streams Supported (SxFy) – H – E, E – H",
    "Communication State Diagram and State Ttransition table (State machine table)",
    "Control States table – offline, on0line/local, online/remote",
    "Equipment Control State transitions",
    "Remote Commands supported – what commands (S2f41 or 49 identify) – with parameters",
    "Error messages table supported (S9F1 – S9F11)",
    "Spooling support4d commands",
    "Data collection from SML",
    "Base SML",
    "From the log files, add to the report table (tool charactereization) and add the the commands to the column",
    "Equpment initiated messages",
    "Reports and events",
    "Remote comma nds",
    "Trace Data (s2f23 – s6f1 –",
    "Recipe managemtn",
    "Imporatne events – s6f11 logs"
]

doc_path = r'e:\Github\EquipmentAutomationPlatforms\Extraction Requirements.docx'

def append_to_doc():
    try:
        print(f"Opening document: {doc_path}")
        doc = Document(doc_path)
        
        # Add some space and a heading before the newly appended items
        doc.add_paragraph("")
        doc.add_heading("Additional Items to Merge", level=2)
        
        for item in items_to_append:
            if item.strip():  # ignore empty strings
                doc.add_paragraph(f"• {item.strip()}")
                
        doc.save(doc_path)
        print("Successfully appended items to the document!")
        
    except Exception as e:
        print(f"An error occurred: {e}")
        print("Please make sure the document is not currently open in Microsoft Word while running this script.")

if __name__ == "__main__":
    append_to_doc()
